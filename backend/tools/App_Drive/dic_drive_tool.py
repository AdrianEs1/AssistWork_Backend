from typing import Dict, Any, Optional, List
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from google.auth.transport.requests import Request
from google.auth.exceptions import RefreshError
from sqlalchemy.orm import Session
from apps.models.oauth_connection import OAuthConnection
from apps.models.conversation import Conversation
from datetime import datetime
import io
import base64
import mimetypes
import fitz  # PyMuPDF (para PDF)
import docx
import os
import re
from backend.apps.services.oauth.google_service_base.google_service_base import GoogleServiceBase

from fastapi import Depends
from apps.core.dependencies import get_db
from apps.database import SessionLocal



# Scopes mínimos para Drive
#SCOPES = ["https://www.googleapis.com/auth/drive"]

class DriveService(GoogleServiceBase):
    def __init__(self):
        super().__init__(service_name="drive", api_version="v3")
    
    def _ping_service(self, service):
        """Verifica conexión con Drive obteniendo info del usuario."""
        about = service.about().get(fields="user").execute()
        user_info = about.get("user", {})
        
        print(f"ℹ️ Ping Drive ejecutado correctamente. Usuario: {user_info}")
        
        return {
            "displayName": user_info.get("displayName"),
            "emailAddress": user_info.get("emailAddress"),
            "permissionId": user_info.get("permissionId"),
            "photoLink": user_info.get("photoLink")
        }


drive = DriveService()


def normalize_name(name: str) -> str:
    """Normaliza nombres para comparación flexible."""
    return re.sub(r"[\s_\-]", "", name.lower())

def sanitize_query(query: str) -> str:
    """
    Limpia queries tipo Drive API:
    name contains 'ClientesOptimusAgent'
    """
    if not query:
        return query

    q = query.lower().strip()

    if "name contains" in q:
        q = q.replace("name contains", "").strip()

    # remover comillas simples o dobles
    q = q.strip("'\"")

    return q


def list_files(
    user_id: str,
    query: Optional[str] = None,
    max_results: int = 10,
    auto_select: bool = True,
    **kwargs
):
    """
    Lista archivos previamente seleccionados por el usuario (Google Drive Picker).

    🔑 Fuente de verdad: Conversation.context["files"]
    🔥 NO consulta Google Drive
    📂 Los archivos son globales al usuario (todas sus conversaciones)
    """
    db= SessionLocal()
    try:
        # 1️⃣ Obtener todas las conversaciones del usuario que tengan contexto
        conversations: List[Conversation] = (
            db.query(Conversation)
            .filter(
                Conversation.user_id == user_id,
                Conversation.context.isnot(None)
            )
            .all()
        )

        # 2️⃣ Recolectar archivos desde el contexto
        files_map = {}  # file_id -> file dict (deduplicación)

        for conv in conversations:
            context = conv.context or {}
            files = context.get("files", [])

            for f in files:
                file_id = f.get("id")
                if file_id and file_id not in files_map:
                    files_map[file_id] = {
                        "id": file_id,
                        "name": f.get("name"),
                        "mimeType": f.get("mime_type"),
                        "source": f.get("source", "google_drive")
                    }

        all_files = list(files_map.values())

        # 3️⃣ No hay archivos seleccionados
        if not all_files:
            return {
                "success": True,
                "files": [],
                "auto_selected": False,
                "needs_user_choice": False,
                "message": (
                    "📎 **No hay archivos disponibles**\n\n"
                    "Primero debes seleccionar archivos desde Google Drive para poder usarlos."
                )
            }

        # 4️⃣ Sin query → mostrar archivos disponibles
        if not query:
            return {
                "success": True,
                "files": all_files[:max_results],
                "auto_selected": False,
                "needs_user_choice": True,
                "message": (
                    f"📂 **Archivos disponibles ({len(all_files)})**\n\n"
                    "Indica el nombre del archivo que deseas usar."
                )
            }

        # 5️⃣ Buscar coincidencias por nombre
        query = sanitize_query(query)
        normalized_query = normalize_name(query)
        matched_files = []

        for f in all_files:
            normalized_name = normalize_name(f["name"])

            if (
                normalized_query == normalized_name
                or normalized_query in normalized_name
            ):
                matched_files.append(f)

        # 6️⃣ No hubo coincidencias
        if not matched_files:
            return {
                "success": True,
                "files": [],
                "auto_selected": False,
                "needs_user_choice": False,
                "message": (
                    "📂 **Archivo no disponible**\n\n"
                    f"No encontré un archivo llamado **{query}** entre los archivos seleccionados.\n\n"
                    "👉 Selecciona el archivo desde Google Drive y vuelve a intentarlo."
                )
            }

        matched_files = matched_files[:max_results]

        # 7️⃣ Un solo archivo → auto select
        if len(matched_files) == 1 and auto_select:
            file = matched_files[0]
            return {
                "success": True,
                "files": matched_files,
                "auto_selected": True,
                "selected_file": file,
                "needs_user_choice": False,
                "message": (
                    "✅ **Archivo seleccionado automáticamente**\n\n"
                    f"📄 **{file['name']}**\n"
                    f"🆔 `{file['id']}`"
                )
            }

        # 8️⃣ Múltiples archivos → pedir confirmación
        user_message = f"🤔 **Encontré {len(matched_files)} archivos. ¿Cuál deseas usar?**\n\n"

        for i, f in enumerate(matched_files, 1):
            user_message += (
                f"**{i}. {f['name']}**\n"
                f"🆔 `{f['id']}`\n\n"
            )

        user_message += "💬 Responde con el número del archivo."

        return {
            "success": True,
            "files": matched_files,
            "auto_selected": False,
            "needs_user_choice": True,
            "message": user_message
        }

    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "auto_selected": False,
            "needs_user_choice": False,
            "message": f"❌ **Error resolviendo archivos**\n\n{str(e)}"
        }
    finally:                     # 4️⃣ SIEMPRE se ejecuta
        db.close()      



"""def list_files(user_id: str, query: Optional[str] = None, max_results: int = 10, auto_select: bool = True, **kwargs):
    ##""Lista archivos con normalización de búsqueda
    try:
        service = drive.get_service(user_id)
        
        all_files = []
        seen_ids = set()
        search_strategies = []
        
        if query:
            if not any(op in query for op in ["name =", "name contains", "mimeType", "in owners"]):
                
                # 🔧 NORMALIZACIÓN: Crear variantes de la query
                query_variants = [
                    query,                           # Original: "tendencias de empleo"
                    query.replace(" ", "_"),         # Con guiones bajos: "tendencias_de_empleo"
                    query.replace(" ", "-"),         # Con guiones: "tendencias-de-empleo"
                    query.replace(" ", ""),          # Sin espacios: "tendenciasdeempleo"
                ]
                
                # Agregar todas las variantes como estrategias de búsqueda
                for variant in query_variants:
                    search_strategies.append({
                        "name": f"variante_{variant[:20]}",
                        "query": f"name contains '{variant}' and trashed = false"
                    })
                
                # También buscar por palabras individuales (fallback)
                words = [w for w in query.split() if len(w) > 2]
                if len(words) > 1:
                    # Probar palabras con diferentes separadores
                    for word in words:
                        search_strategies.append({
                            "name": f"palabra_{word}",
                            "query": f"name contains '{word}' and trashed = false"
                        })
            else:
                # Query personalizada con operadores Drive
                search_strategies.append({
                    "name": "custom",
                    "query": f"({query}) and trashed = false"
                })
        else:
            # Sin query: archivos recientes
            search_strategies.append({
                "name": "recientes",
                "query": "trashed = false"
            })
        
        # Ejecutar estrategias hasta encontrar resultados suficientes
        for strategy in search_strategies:
            try:
                results = service.files().list(
                    q=strategy["query"],
                    pageSize=max_results,
                    fields="files(id, name, mimeType, modifiedTime, size, webViewLink, owners)",
                    orderBy="modifiedTime desc"
                ).execute()
                
                for file in results.get("files", []):
                    if file['id'] not in seen_ids:
                        all_files.append(file)
                        seen_ids.add(file['id'])
                
                # Si encontramos al menos 1 archivo, continuamos buscando
                # pero si ya tenemos suficientes, paramos
                if len(all_files) >= max_results:
                    print(f"✅ Suficientes resultados con '{strategy['name']}'")
                    break
                    
            except Exception as e:
                print(f"⚠️ Estrategia '{strategy['name']}' falló: {e}")
                continue
        
        # Si no encontramos nada, retornar mensaje claro
        if not all_files:
            return {
                "success": True,
                "files": [],
                "auto_selected": False,
                "needs_user_choice": False,
                "message": f"📂 **No se encontraron archivos**\n\nNo hay archivos que coincidan con: `{query}`\n\n💡 Intenta con términos más generales o verifica el nombre exacto."
            }
        
        # Formatear archivos
        detailed_files = []
        for f in all_files[:max_results]:
            detailed_files.append({
                "id": f["id"],
                "name": f["name"],
                "mimeType": f.get("mimeType"),
                "modifiedTime": f.get("modifiedTime"),
                "size": f.get("size"),
                "webViewLink": f.get("webViewLink"),
                "owners": f.get("owners", [])
            })
        
        # CASO 1: Solo 1 archivo → Selección automática
        if len(detailed_files) == 1 and auto_select:
            file = detailed_files[0]
            return {
                "success": True,
                "files": detailed_files,
                "auto_selected": True,
                "selected_file": file,
                "needs_user_choice": False,
                "message": f"✅ **Archivo encontrado automáticamente**\n\n📄 **{file['name']}**\n🆔 ID: `{file['id']}`\n📅 Modificado: {file.get('modifiedTime', 'N/A')[:10]}"
            }
        
        # CASO 2: Múltiples archivos (2-5) → Pedir confirmación
        elif 2 <= len(detailed_files) <= 5:
            user_message = f"🤔 **Encontré {len(detailed_files)} archivos. ¿Cuál prefieres?**\n\n"
            
            for i, f in enumerate(detailed_files, 1):
                user_message += f"**{i}. {f['name']}**\n"
                user_message += f"   📁 Tipo: {f.get('mimeType', 'Desconocido')}\n"
                user_message += f"   🆔 ID: `{f['id']}`\n"
                user_message += f"   📅 Modificado: {f.get('modifiedTime', 'N/A')[:10]}\n\n"
            
            user_message += "💬 **Responde con el número del archivo que quieres usar.**"
            
            return {
                "success": True,
                "files": detailed_files,
                "auto_selected": False,
                "needs_user_choice": True,
                "message": user_message
            }
        
        # CASO 3: Muchos archivos (>5) → Mostrar top 10
        else:
            user_message = f"📋 **Encontré {len(detailed_files)} archivos. Mostrando los más relevantes:**\n\n"
            
            for i, f in enumerate(detailed_files[:10], 1):
                user_message += f"**{i}. {f['name']}**\n"
                user_message += f"   📅 {f.get('modifiedTime', 'N/A')[:10]}\n"
                user_message += f"   🆔 `{f['id']}`\n\n"
            
            user_message += "💬 **Opciones:**\n"
            user_message += "• Responde con el **número** del archivo\n"
            user_message += "• O dame más detalles para refinar la búsqueda"
            
            return {
                "success": True,
                "files": detailed_files[:10],
                "auto_selected": False,
                "needs_user_choice": True,
                "message": user_message
            }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "auto_selected": False,
            "needs_user_choice": False,
            "message": f"❌ **Error listando archivos**\n\n🚫 {str(e)}"
        }"""

# ---------------------------------------------
# 📖 Leer contenido de un archivo
# ---------------------------------------------


def read_file(user_id: str, file_id: str, **kwargs):
    """
    Lee el contenido completo de un archivo en Google Drive (PDF, DOCX, TXT, etc.).
    Retorna en formato estandarizado compatible con el orquestador.
    """
    try:
        service = drive.get_service(user_id)
        
        # Si el argumento parece un nombre en lugar de un ID, buscarlo
        if len(file_id) < 20 and not any(c in file_id for c in "-_"):
            print(f"🔍 Buscando archivo por nombre: {file_id}")
            query = f"name contains '{file_id}' and trashed = false"
            results = service.files().list(
                q=query,
                pageSize=1,
                fields="files(id, name, mimeType)"
            ).execute()
            files = results.get("files", [])
            if not files:
                return {"success": False, "message": f"❌ No se encontró ningún archivo con nombre '{file_id}'."}
            file_id = files[0]["id"]
            print(f"✅ Archivo encontrado: {files[0]['name']} → ID: {file_id}")
        
        # Obtener metadatos
        file = service.files().get(fileId=file_id, fields="name, mimeType").execute()
        file_name = file["name"]
        mime_type = file["mimeType"]
        
        # Descargar contenido binario
        request = service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        fh.seek(0)
        
        # Procesar texto según tipo MIME
        content = ""
        if mime_type == "application/pdf":
            import fitz
            with fitz.open(stream=fh, filetype="pdf") as doc:
                content = "\n".join(page.get_text("text") for page in doc)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            from docx import Document
            doc = Document(fh)
            content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        else:
            content = fh.read().decode(errors="ignore")
        
        # Configuración de preview (similar a Gmail)
        preview_length = kwargs.get('preview_length', 500)
        show_full = kwargs.get('show_full', False)
        
        # Mensaje formateado para el usuario
        user_message = f"📄 **Archivo leído exitosamente**\n\n"
        user_message += f"📝 **Nombre:** {file_name}\n"
        user_message += f"📋 **Tipo:** {mime_type}\n"
        user_message += f"📊 **Tamaño:** {len(content)} caracteres\n"
        user_message += f"🆔 **ID:** `{file_id}`\n\n"
        
        if show_full or len(content) <= preview_length:
            user_message += f"📄 **Contenido:**\n```\n{content}\n```"
        else:
            user_message += f"📄 **Contenido:**\n```\n{content[:preview_length]}...\n```\n\n*(Vista previa truncada)*"
        
        # Retorno estandarizado (mismo formato que Gmail)
        return {
            "success": True,
            "file_id": file_id,
            "file_name": file_name,
            "mime_type": mime_type,
            "content": content,  # ✅ Contenido completo en primer nivel
            "content_length": len(content),
            "content_preview": content[:preview_length] if len(content) > preview_length else content,
            "message": user_message
        }
       
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "message": f"❌ **Error leyendo archivo**\n\n🆔 **ID:** `{file_id}`\n🚫 **Error:** {str(e)}"
        }


# ---------------------------------------------
# ⬆️ Subir archivo
# ---------------------------------------------
def upload_file(user_id: str, file_path: str, mime_type: Optional[str] = None, **kwargs) -> Dict[str, Any]:
    try:
        service = drive.get_service(user_id)
        if not mime_type:
            mime_type, _ = mimetypes.guess_type(file_path)

        file_metadata = {"name": os.path.basename(file_path)}
        media = MediaFileUpload(file_path, mimetype=mime_type)
        uploaded = service.files().create(body=file_metadata, media_body=media, fields="id, name").execute()

        return {
            "success": True,
            "file_id": uploaded["id"],
            "file_name": uploaded["name"],
            "message": f"✅ Archivo **{uploaded['name']}** subido exitosamente a Google Drive."
        }

    except Exception as e:
        return {"success": False, "error": str(e), "message": f"❌ Error subiendo archivo: {str(e)}"}


# ---------------------------------------------
# 🗑️ Eliminar archivo
# ---------------------------------------------
def delete_file(user_id: str, file_id: str, **kwargs) -> Dict[str, Any]:
    try:
        service = drive.get_service(user_id)
        service.files().delete(fileId=file_id).execute()
        return {"success": True, "message": "🗑️ Archivo eliminado correctamente."}
    except Exception as e:
        return {"success": False, "error": str(e), "message": f"❌ Error eliminando archivo: {str(e)}"}


# ---------------------------------------------
# 🔍 Probar conexión
# ---------------------------------------------
def test_connection(user_id: str, **kwargs) -> Dict[str, Any]:
    connect_service = drive.test_connection(user_id)
    return  connect_service 


# ---------------------------------------------
# 🧩 Auxiliares para lectura de texto
# ---------------------------------------------
def _extract_text_from_pdf(fh: io.BytesIO) -> str:
    text = ""
    with fitz.open(stream=fh.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text("text") + "\n"
    return text.strip()


def _extract_text_from_docx(fh: io.BytesIO) -> str:
    text = ""
    with open("temp.docx", "wb") as f:
        f.write(fh.read())
    doc = docx.Document("temp.docx")
    for para in doc.paragraphs:
        text += para.text + "\n"
    os.remove("temp.docx")
    return text.strip()


# ---------------------------------------------
# 📚 Registro para el agente
# ---------------------------------------------
DRIVE_TOOL_METHODS = {
    "list_files": {"func": list_files, "description": "Listar archivos en Google Drive del usuario."},
    "read_file": {"func": read_file, "description": "Leer y devolver el contenido textual de un archivo de Google Drive (PDF, DOCX, TXT, etc.)."},
    "upload_file": {"func": upload_file, "description": "Subir un archivo local o generado por el agente a Google Drive."},
    "delete_file": {"func": delete_file, "description": "Eliminar un archivo del Drive del usuario."},
    "test_connection": {"func": test_connection, "description": "Probar si la conexión a Google Drive está activa y autenticada."}
}
